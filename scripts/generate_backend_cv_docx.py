"""
Generate a Word (.docx) version of `docs/backend_cv.md`.
Requires: python-docx (install with `pip install python-docx`)
Run: python scripts/generate_backend_cv_docx.py
"""
from pathlib import Path
from docx import Document

src = Path(__file__).resolve().parents[1] / 'docs' / 'backend_cv.md'
if not src.exists():
    raise SystemExit('docs/backend_cv.md not found')

text = src.read_text(encoding='utf-8')
lines = text.splitlines()

doc = Document()
doc.add_heading('Backend CV — AI Cyber Threat Prevention', level=1)

p = None
for line in lines:
    if line.startswith('# '):
        # skip main heading (already added)
        continue
    if line.startswith('## '):
        doc.add_heading(line[3:].strip(), level=2)
        p = None
        continue
    if line.startswith('```'):
        # simple code block handling
        code_lines = []
        i = lines.index(line)
        # collect until next ```
        j = i + 1
        while j < len(lines) and not lines[j].startswith('```'):
            code_lines.append(lines[j])
            j += 1
        doc.add_paragraph('\n'.join(code_lines), style='Courier')
        # move the outer loop forward (simple approach)
        # (not a full parser; for this small doc it suffices)
        break
    if line.strip() == '':
        p = None
        continue
    # normal paragraph
    doc.add_paragraph(line)

out = src.with_suffix('.docx')
doc.save(out)
print('Wrote', out)
